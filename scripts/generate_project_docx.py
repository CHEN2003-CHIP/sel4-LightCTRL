#!/usr/bin/env python3
"""Generate simple DOCX files from project-mds Markdown files.

The converter intentionally keeps Mermaid blocks as code when no image
renderer is available. It is designed for reviewable course documents, not as
a full Markdown implementation.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SRC_DIR = ROOT / "project-mds"
OUT_DIR = ROOT / "project-docs"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def normalize_filename(path: Path) -> str:
    stem = path.stem
    if stem == "2.概要设计":
        stem = "2.概要设计说明书"
    return f"{stem}.docx"


def apply_styles(document: Document) -> None:
    styles = document.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for name, size in [("Heading 1", 18), ("Heading 2", 15), ("Heading 3", 13), ("Heading 4", 12)]:
        style = styles[name]
        style.font.name = "黑体"
        style.font.size = Pt(size)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_code_block(document: Document, lines: list[str], language: str) -> None:
    title = "代码块"
    if language:
        title = f"代码块：{language}"
    paragraph = document.add_paragraph(title)
    paragraph.style = "Intense Quote"
    for line in lines:
        p = document.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Pt(18)


def is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def flush_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    filtered = []
    for row in rows:
        if row and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in row):
            continue
        filtered.append(row)
    if not filtered:
        return

    cols = max(len(row) for row in filtered)
    table = document.add_table(rows=len(filtered), cols=cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(filtered):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.text = row[c_idx] if c_idx < len(row) else ""
            if r_idx == 0:
                set_cell_shading(cell, "D9EAF7")


def add_markdown_line(document: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        document.add_paragraph("")
        return

    match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
    if match:
        level = min(len(match.group(1)), 4)
        document.add_heading(match.group(2), level=level)
        return

    if stripped.startswith("- "):
        document.add_paragraph(stripped[2:], style="List Bullet")
        return

    ordered = re.match(r"^\d+\.\s+(.*)$", stripped)
    if ordered:
        document.add_paragraph(ordered.group(1), style="List Number")
        return

    paragraph = document.add_paragraph()
    # Basic bold marker cleanup for document readability.
    cleaned = stripped.replace("**", "")
    paragraph.add_run(cleaned)


def convert_one(src: Path, dst: Path) -> None:
    document = Document()
    apply_styles(document)

    section = document.sections[0]
    section.header.paragraphs[0].text = "LightDemo 软件工程文档"
    section.footer.paragraphs[0].text = "由 project-mds 生成；Mermaid 图如未渲染则保留为代码块。"

    lines = src.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    table_rows: list[list[str]] = []

    for line in lines:
        if line.startswith("```"):
            if in_code:
                add_code_block(document, code_lines, code_lang)
                in_code = False
                code_lang = ""
                code_lines = []
            else:
                flush_table(document, table_rows)
                table_rows = []
                in_code = True
                code_lang = line.strip().strip("`").strip()
            continue

        if in_code:
            code_lines.append(line)
            continue

        if is_table_line(line):
            table_rows.append(split_table_row(line))
            continue

        flush_table(document, table_rows)
        table_rows = []
        add_markdown_line(document, line)

    if in_code:
        add_code_block(document, code_lines, code_lang)
    flush_table(document, table_rows)

    # Add a final page break marker friendly to Word editing.
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.save(dst)


def main() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    markdown_files = sorted(SRC_DIR.glob("*.md"), key=lambda p: p.name)
    for src in markdown_files:
        dst = OUT_DIR / normalize_filename(src)
        convert_one(src, dst)
        print(f"generated {dst}")


if __name__ == "__main__":
    main()
